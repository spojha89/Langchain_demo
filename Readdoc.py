from docx import Document
# import streamlit as st

# text=[]
doc_path='C:\\Users\\Reshma\\Downloads\\Writing_flight_miss_Subhra.docx'


# doc_reader= Document(doc)
# for para in doc_reader.paragraphs:
#     text.append(para.text)

# print(text)

# doc_path='C:\\Users\\Reshma\\Downloads\\Cover_Letter_Subhra_Ojha.docx'

doc2=Document(doc_path)

text_box_texts = []

for para in doc2.paragraphs:
    text_box_texts.append(para.text)

for shape in doc2.inline_shapes:
    if shape.type==3:
        if shape.text_frame:
            for paragraph in shape.text_frame.paragraphs:
                text_box_texts.append(paragraph.text)

print(text_box_texts)